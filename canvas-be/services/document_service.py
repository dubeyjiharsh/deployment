from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
import tempfile
import os
import json
import logging

class DocumentService:
    """Service for generating DOCX and PDF documents from canvas data"""

    @staticmethod
    def _get_clean_data():
        """
        Returns configuration for document generation, including fields to exclude and fields to render as tables.
        Excludes only normalized versions of canvas_id, id, created_at, updated_at, and tags.
        Returns:
            dict: Configuration with 'exclude' set and 'tables' set.
        """
        exclude = {k.replace("_", " ").lower().strip() for k in [
            "canvas_id", "id", "created_at", "updated_at", "tags"
        ]}
        return {
            "exclude": exclude,
            "tables": {
                "kpis",
                "key features",
                "risks",
                "non functional requirements",
                "use cases"
            }
        }

    @staticmethod
    def _normalize_table_data(value):
        """
        Normalize various input formats (list of dicts or list of JSON strings) into a list of dicts.
        Args:
            value (list): List of dicts or JSON strings.
        Returns:
            list: List of dicts after normalization.
        """
        if not isinstance(value, list):
            return []
        normalized = []
        for item in value:
            if isinstance(item, dict):
                normalized.append(item)
            elif isinstance(item, str):
                try:
                    # Handle cases where the string might be wrapped in extra quotes or whitespace
                    cleaned_item = item.strip()
                    normalized.append(json.loads(cleaned_item))
                except (json.JSONDecodeError, TypeError):
                    continue
        return normalized

    @staticmethod
    def generate_docx(canvas_data: dict) -> str:
        """
        Generate a DOCX document from the provided canvas data.
        Args:
            canvas_data (dict): The canvas data to render.
        Returns:
            str: Path to the generated DOCX file.
        """
        logger = logging.getLogger("DocumentService")
        doc = Document()
        config = DocumentService._get_clean_data()
        table_keys = config["tables"]

        # Title
        doc.add_heading(canvas_data.get("Title", "Business Model Canvas"), 0)

        for key, value in canvas_data.items():
            key_clean = key.replace("_", " ").lower().strip()
            # Only skip normalized exclude fields, but keep relevant facts
            if key == "Title" or key_clean in config["exclude"] or not value:
                continue

            doc.add_heading(key.replace("_", " ").title(), level=1)

            # TABLE LOGIC (KPIs, Key Features, Risks, NFRs, Use Cases)
            if key_clean in table_keys and isinstance(value, list):
                normalized = DocumentService._normalize_table_data(value)
                
                if normalized and isinstance(normalized[0], dict):
                    columns = list(normalized[0].keys())
                    table = doc.add_table(rows=1, cols=len(columns))
                    table.style = "Table Grid"
                    
                    # Header row
                    for i, col in enumerate(columns):
                        table.rows[0].cells[i].text = col.replace("_", " ").title()
                    
                    # Data rows
                    for entry in normalized:
                        row_cells = table.add_row().cells
                        for i, col in enumerate(columns):
                            row_cells[i].text = str(entry.get(col, ""))
                else:
                    # Fallback to bullet points if data isn't structured
                    for item in value:
                        doc.add_paragraph(str(item), style="List Bullet")

            # LIST FIELDS (Simple strings in a list)
            elif isinstance(value, list):
                for item in value:
                    doc.add_paragraph(str(item), style="List Bullet")

            # DICT FIELDS (Key-Value pairs)
            elif isinstance(value, dict):
                for k, v in value.items():
                    p = doc.add_paragraph()
                    p.add_run(f"{k.replace('_', ' ').title()}: ").bold = True
                    p.add_run(str(v))

            # PLAIN TEXT 
            else:
                doc.add_paragraph(str(value))

        temp_file = os.path.join(tempfile.gettempdir(), f"canvas_{os.urandom(4).hex()}.docx")
        doc.save(temp_file)
        return temp_file

    @staticmethod
    def generate_pdf(canvas_data: dict) -> str:
        """
        Generate a PDF document from the provided canvas data.
        Args:
            canvas_data (dict): The canvas data to render.
        Returns:
            str: Path to the generated PDF file.
        """
        logger = logging.getLogger("DocumentService")
        config = DocumentService._get_clean_data()
        table_keys = config["tables"]

        temp_file = os.path.join(tempfile.gettempdir(), f"canvas_{os.urandom(4).hex()}.pdf")
        doc = SimpleDocTemplate(temp_file, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []

        # Title
        elements.append(Paragraph(canvas_data.get("Title", "Business Model Canvas"), styles["Title"]))
        elements.append(Spacer(1, 12))

        for key, value in canvas_data.items():
            key_clean = key.replace("_", " ").lower().strip()
            # Only skip normalized exclude fields, but keep relevant facts
            if key == "Title" or key_clean in config["exclude"] or not value:
                continue

            elements.append(Paragraph(key.replace("_", " ").title(), styles["Heading1"]))
            elements.append(Spacer(1, 6))

            # TABLE LOGIC
            if key_clean in table_keys and isinstance(value, list):
                normalized = DocumentService._normalize_table_data(value)
                
                if normalized and isinstance(normalized[0], dict):
                    headers = [h.replace("_", " ").title() for h in normalized[0].keys()]
                    data = [headers]
                    
                    for entry in normalized:
                        row = [Paragraph(str(entry.get(k, "")), styles["BodyText"]) for k in normalized[0].keys()]
                        data.append(row)
                    
                    # Calculate column widths evenly
                    col_widths = [doc.width / len(headers)] * len(headers)
                    table = Table(data, colWidths=col_widths)
                    table.setStyle(TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.beige]),
                    ]))
                    elements.append(table)
                else:
                    for item in value:
                        elements.append(Paragraph(f"• {item}", styles["BodyText"]))
                elements.append(Spacer(1, 12))

            # LIST FIELDS
            elif isinstance(value, list):
                for item in value:
                    elements.append(Paragraph(f"• {item}", styles["BodyText"]))
                elements.append(Spacer(1, 12))

            # DICT FIELDS
            elif isinstance(value, dict):
                for k, v in value.items():
                    elements.append(Paragraph(f"<b>{k.replace('_', ' ').title()}:</b> {v}", styles["BodyText"]))
                elements.append(Spacer(1, 12))

            # TEXT
            else:
                elements.append(Paragraph(str(value), styles["BodyText"]))
                elements.append(Spacer(1, 12))

        doc.build(elements)
        return temp_file